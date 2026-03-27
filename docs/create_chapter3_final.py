"""
졸업설계 보고서 — 3장 방법론 (간결판)
표 10개, 식 4개. 개조식 한국어 학술체.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pathlib

OUTPUT_DIR = pathlib.Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    shading.append(elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='2E4057'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def caption(doc, num, title):
    p = doc.add_paragraph()
    p.add_run(f'표 3-{num}. ').bold = True
    p.add_run(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def eq(doc, eq_num, eq_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(eq_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(f'({eq_num})')


def body(doc, text):
    doc.add_paragraph(text)


def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def bullets(doc, items):
    for item in items:
        bullet(doc, item)


def create_report():
    doc = Document()

    # 스타일
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.name = '맑은 고딕'

    # =================================================================
    doc.add_heading('3. 방법론', level=1)
    # =================================================================

    # -----------------------------------------------------------------
    # 3.1 시뮬레이션 프레임워크
    # -----------------------------------------------------------------
    doc.add_heading('3.1 시뮬레이션 프레임워크', level=2)

    body(doc,
        '태그리스 단계적 도입 과정에서, 태그리스 하드웨어가 설치된 게이트를 '
        '겸용으로 운영할 것인가 전용(express lane)으로 운영할 것인가에 따른 '
        '통행비용 차이를 미시적 보행 시뮬레이션으로 분석한다.')

    add_styled_table(doc,
        ['구분', '내용'],
        [
            ['시뮬레이션 도구', 'JuPedSim (Python/C++, 오픈소스, LGPL)'],
            ['보행 모델', 'CFSM V2 (Tordeux et al., 2016; Xu et al., 2019)'],
            ['게이트 선택', 'MNL — Gao et al. (2019) LRP 기반'],
            ['대상역', '성수역 2호선 서쪽 대합실 (50m × 25m, 게이트 7개)'],
            ['V&V', 'NIST TN 1822 (Ronchi et al., 2013) 준용'],
        ],
        col_widths=[5, 12]
    )
    caption(doc, 1, '시뮬레이션 프레임워크 구성')

    doc.add_heading('3.1.1 도구 선정 근거', level=3)

    body(doc, 'JuPedSim 선정 근거:')
    bullets(doc, [
        'CFSM V2 기본 지원 — 0.55m 게이트 통로에서 안정적 시뮬레이션',
        '에이전트별 파라미터 개별 설정·동적 변경 가능 — 태그/태그리스 서비스 시간 차별화',
        'Python API — MNL 게이트 선택, Choice Set 분리 등 자유 구현',
        '오픈소스 — 코드·파라미터 완전 공개, 재현성 보장',
    ])

    doc.add_heading('3.1.2 AnyLogic 미채택 근거', level=3)

    add_styled_table(doc,
        ['항목', 'JuPedSim', 'AnyLogic'],
        [
            ['보행 모델', 'CFSM V2 (속도 기반, 충돌 방지)', 'SFM (힘 기반, 겹침 위험)'],
            ['게이트 선택', 'MNL 자유 구현', 'PedSelectOutput — 고정확률/조건분기만'],
            ['Choice Set 분리', 'Python 자유 구현', '미지원'],
            ['V&V dt/겹침 검증', '완전 가능', '불가 (블랙박스)'],
            ['코드 재현성', '완전 공개', '바이너리 모델 파일'],
        ],
        col_widths=[4.5, 6.5, 6]
    )
    caption(doc, 2, 'JuPedSim vs AnyLogic 비교')

    body(doc, 'AnyLogic PedSelectOutput의 한계:')
    bullets(doc, [
        'Probabilities — 고정 비율 분배, 대기열·거리에 따른 동적 반응 없음',
        'Conditions — Boolean 조건 순차 분기, 확률적 선택 불가',
        'Exit Number — 결정론적 출구 지정, 확률적 선택 불가',
    ])

    # -----------------------------------------------------------------
    # 3.2 분석 대상
    # -----------------------------------------------------------------
    doc.add_heading('3.2 분석 대상: 성수역 서쪽 대합실', level=2)

    body(doc, '대상역 선정 근거:')
    bullets(doc, [
        '환승 없음 → 역사 구조 단순, 시뮬레이션 구현 용이',
        '승객/게이트 비율 상위 3위 (일평균 105,574명 / 39대 = 2,707명/게이트)',
        '태그리스 미설치 → 우이신설선 실측값 이식(transferability) 방식 적용',
    ])

    add_styled_table(doc,
        ['구성요소', '치수/수량', '비고'],
        [
            ['대합실', '50m × 25m', '상부 노치 포함'],
            ['계단', '2개, 폭 3.0m', '상부(y=15~18), 하부(y=8~11)'],
            ['출구', '2개, 폭 3.0m', '상부(y=24), 하부(y=3)'],
            ['게이트', '7개', '통로 폭 0.55m, 하우징 0.30m'],
            ['게이트 배리어', 'x=12.0m, y=3~22m', '클러스터 높이 6.25m'],
        ],
        col_widths=[4, 5, 8]
    )
    caption(doc, 3, '대합실 기하구조 및 게이트 배치')

    # -----------------------------------------------------------------
    # 3.3 보행 모델
    # -----------------------------------------------------------------
    doc.add_heading('3.3 보행 모델: CFSM V2', level=2)

    body(doc, 'SFM 대비 CFSM V2 선정 근거:')
    bullets(doc, [
        'SFM: 힘 기반(F=ma) → 좁은 통로에서 겹침·진동 발생 (Kretz, 2015)',
        'CFSM V2: 1차 속도 방정식 → 간격이 직경 이하이면 속도=0, 충돌 구조적 방지',
        '초기 적용한 GCFM은 V5(dt 변동 24%)·V6(겹침 22.4cm) FAIL → CFSM V2로 전환',
    ])

    body(doc, '수학적 정식화:')

    eq(doc, 1, 'ẋᵢ = V(sᵢ) · eᵢ')
    eq(doc, 2, 'V(s) = min{ v₀, max{ 0, (s − l) / T } }')

    bullets(doc, [
        'V(s): 최적 속도 함수. s ≤ l이면 V=0 → 겹침 원천 방지',
        'v₀: 희망속도, l: 보행자 직경(=2r), T: 시간 간격, s: 전방 이웃 거리',
        '이동 방향 eᵢ: 목적지 방향 + 이웃·벽면 반발의 합',
    ])

    add_styled_table(doc,
        ['파라미터', '기호', '값', '출처'],
        [
            ['시간 간격', 'T', '0.80 s', 'Seyfried (2009) 캘리브레이션'],
            ['보행자 반경', 'r', '0.15 m', '0.55m 게이트 통과 대응'],
            ['이웃 반발 강도/범위', 'aₙ / Dₙ', '8.0 / 0.1 m', 'Tordeux et al. (2016)'],
            ['벽면 반발 강도/범위', 'ag / Dg', '5.0 / 0.02 m', 'Tordeux et al. (2016)'],
            ['희망속도', 'v₀', 'N(1.34, 0.26) m/s', 'Weidmann (1993)'],
            ['시간 스텝', 'dt', '0.05 s', '수치 안정성'],
        ],
        col_widths=[4, 2.5, 5.5, 5]
    )
    caption(doc, 4, 'CFSM V2 파라미터')

    # -----------------------------------------------------------------
    # 3.4 게이트 선택 및 서비스 시간 모델
    # -----------------------------------------------------------------
    doc.add_heading('3.4 게이트 선택 및 서비스 시간 모델', level=2)

    doc.add_heading('3.4.1 MNL 게이트 선택', level=3)

    body(doc, 'Gao et al. (2019) LRP 기반 다항 로짓 모형:')

    eq(doc, 3, 'Cⱼ = ωᴺ · Wⱼ + ωᴸ · (L₁ⱼ + L₃ⱼ) / vᵢ')
    eq(doc, 4, 'P(j) = exp(−Cⱼ) / Σₖ exp(−Cₖ),  k ∈ Choice Set')

    bullets(doc, [
        'Cⱼ: 게이트 j 비용, Wⱼ: 예상 대기시간, L₁ⱼ: 접근거리, L₃ⱼ: 게이트→출구 거리',
        '거리 추정: 순서보존 노이즈 (Gao eq.4-5), 대기열 인지: 인지오차 (Gao eq.7)',
        '3단계 의사결정: 3.0m(MNL) → 1.7m(MNL 재평가) → 1.0m(인접 빈 게이트 확정 전환)',
        '핑퐁 방지: 관성(C_switch=1.5), lock-in(3.0m), 재평가 주기(3.0s)',
    ])

    add_styled_table(doc,
        ['성격 유형', 'ωᴺ (대기)', 'ωᴸ (보행)', '행태', '비율'],
        [
            ['adventurous', '1.2', '0.8', '대기 회피, 먼 빈 게이트 감수', '1/3'],
            ['conserved', '0.8', '1.2', '가까운 게이트 선호', '1/3'],
            ['mild', '1.0', '1.0', '균등 가중', '1/3'],
        ],
        col_widths=[3, 2.5, 2.5, 6, 2]
    )
    caption(doc, 5, '보행자 성격 유형별 가중치 (Gao et al., 2019)')

    doc.add_heading('3.4.2 이용자 유형별 Choice Set (연구 핵심)', level=3)

    body(doc, '이용자 유형에 따라 선택 가능 게이트 집합을 분리한다:')

    add_styled_table(doc,
        ['시나리오', '태그 이용자', '태그리스 이용자'],
        [
            ['기본 (겸용 운영)', '태그전용 + 겸용 게이트', '겸용 게이트만'],
            ['개선 (express lane)', '태그 전용 게이트만', '태그리스 전용 게이트만'],
        ],
        col_widths=[4.5, 6, 6.5]
    )
    caption(doc, 6, '시나리오별 Choice Set')

    bullets(doc, [
        '기본: 겸용 게이트에서 태그/태그리스 혼재 → 태그 이용자 정지·탭으로 태그리스 대기 발생',
        '개선: 완전 분리 → 혼재 비효율 제거',
    ])

    doc.add_heading('3.4.3 서비스 시간', level=3)

    add_styled_table(doc,
        ['이용자', '분포', '파라미터', '비고'],
        [
            ['태그', 'LogNormal', 'μ_ln=0.568, σ_ln=0.5, [0.8, 3.7]s', '평균 2.0s, 통과속도 0.65m/s'],
            ['태그리스', '—', '0s (민감도: 0/0.5/1.0s)', '우이신설선 실측 예정'],
        ],
        col_widths=[3, 3.5, 6.5, 4]
    )
    caption(doc, 7, '서비스 시간 모델')

    bullets(doc, [
        '태그: 게이트 진입 시 점유 상태 전환 → 서비스 시간 동안 0.65m/s로 감속 → 해제',
        '태그리스: 점유 없이 희망속도 유지, 무정차 연속 통과',
    ])

    # -----------------------------------------------------------------
    # 3.5 수요 모델
    # -----------------------------------------------------------------
    doc.add_heading('3.5 수요 모델', level=2)

    add_styled_table(doc,
        ['파라미터', '값', '근거'],
        [
            ['열차 도착 간격', '180 s', '2호선 피크 배차 간격'],
            ['1회 하차 인원', 'Poisson(λ=40)', '서쪽 계단 이용분 추정'],
            ['계단 진입 시각', 'N(7.5, 3.75) s', '열차 도착 시점 기준'],
            ['계단 배분', '상부:하부 = 50:50', '균등 분배'],
            ['시뮬레이션 시간', '330 s', '열차 2회 도착 + 잔류 소화'],
        ],
        col_widths=[5, 5, 7]
    )
    caption(doc, 8, '수요 모델 파라미터')

    # -----------------------------------------------------------------
    # 3.6 시나리오 설계
    # -----------------------------------------------------------------
    doc.add_heading('3.6 시나리오 설계', level=2)

    body(doc, '기본 시나리오 (현행 겸용 운영):')
    bullets(doc, [
        '태그리스 하드웨어 설치 게이트 K대를 겸용(태그+태그리스)으로 운영',
        '나머지(7−K)대는 태그 전용 (하드웨어 미설치)',
        '겸용 게이트에서 태그 이용자 정지·탭 → 태그리스 이용자 대기 발생',
    ])

    body(doc, '개선 시나리오 (express lane 전용 운영):')
    bullets(doc, [
        '동일 K대를 태그리스 전용(express lane)으로 전환',
        '태그리스 이용자 → 전용 게이트만 이용, 무정차 연속 통과',
        '추가 하드웨어 없이 운영 정책(안내판/SW)만 변경 — 하이패스 전용 차로와 동일 개념',
    ])

    add_styled_table(doc,
        ['요인', '수준', '값'],
        [
            ['운영 방식', '2', '겸용 / express lane 분리'],
            ['태그리스 이용자 비율 (p)', '5', '10%, 20%, 40%, 60%, 80%'],
            ['태그리스 게이트 수 (K)', '3', '1, 2, 3대'],
            ['시간대', '2', '피크 / 비피크'],
        ],
        col_widths=[5.5, 2.5, 9]
    )
    caption(doc, 9, '실험 요인 (2×5×3×2 = 60개 시나리오)')

    body(doc,
        '각 시나리오를 N회 반복하여 통계적 유의성 확보. '
        'N은 U1 수렴성 분석(ERD < 5%, 3회 연속)으로 결정.')

    # -----------------------------------------------------------------
    # 3.7 V&V
    # -----------------------------------------------------------------
    doc.add_heading('3.7 V&V (NIST TN 1822)', level=2)

    body(doc, 'NIST TN 1822 6단계 프레임워크 준용:')
    bullets(doc, [
        'Phase 1 — Model Qualification: CFSM V2 + MNL 선정 근거 문서화 (4/5 완료)',
        'Phase 2 — Verification: RiMEA 기반 7개 테스트, 6/7 PASS (V5 재측정 예정)',
        'Phase 3 — Calibration: time_gap 1.06→0.80s 완료(C1), 나머지 현장조사 후',
        'Phase 4 — Validation: 우이신설선 비피크 독립 데이터로 비교',
        'Phase 5 — Uncertainty: (U1) 반복 수렴성, (U2) 9개 파라미터 OAT ±20%, (U3) 시나리오 강건성',
    ])

    # -----------------------------------------------------------------
    # 3.8 평가 지표
    # -----------------------------------------------------------------
    doc.add_heading('3.8 평가 지표', level=2)

    add_styled_table(doc,
        ['지표', '정의', '단위'],
        [
            ['평균 대기시간', '게이트 도착 ~ 서비스 시작', '초'],
            ['최대 대기행렬', '게이트 앞 최대 대기 인원', '명'],
            ['총 통과시간', '계단 출발 ~ 게이트 통과 완료', '초'],
            ['게이트 균형도 (MD)', '게이트별 처리량 평균편차', '%'],
            ['보행자 밀도 / LOS', '게이트-계단 구간 밀도', '인/m²'],
            ['처리량', '단위시간당 게이트 통과 인원', 'person/min'],
        ],
        col_widths=[4.5, 8.5, 4]
    )
    caption(doc, 10, '평가 지표')

    # 저장
    out_path = OUTPUT_DIR / "3장_방법론_v3.docx"
    doc.save(str(out_path))
    print(f"저장 완료: {out_path}")


if __name__ == '__main__':
    create_report()
