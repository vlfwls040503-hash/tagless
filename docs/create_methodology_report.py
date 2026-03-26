"""
졸업설계 보고서 — 3장 방법론 (docx)
V&V 프레임워크 중심, 계획 형태로 작성
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pathlib

OUTPUT_DIR = pathlib.Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='2E4057'):
    """스타일된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 간격
    return table


def create_report():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '맑은 고딕'

    # =========================================================================
    # 3장 제목
    # =========================================================================
    h = doc.add_heading('3. 방법론', level=1)

    # =========================================================================
    # 3.1 시뮬레이션 프레임워크 개요
    # =========================================================================
    doc.add_heading('3.1 시뮬레이션 프레임워크 개요', level=2)

    doc.add_paragraph(
        '본 연구는 태그리스 전용 게이트 분리 배치의 통행비용 절감 효과를 분석하기 위해 '
        '미시적 보행자 시뮬레이션을 수행한다. 시뮬레이션 프레임워크는 세 가지 핵심 모델로 구성된다: '
        '(1) 보행 역학을 담당하는 Collision-Free Speed Model V2 (CFSM V2), '
        '(2) 게이트 선택 의사결정을 담당하는 다항 로짓 모델 (MNL), '
        '(3) 게이트 서비스 시간 모델. '
        '시뮬레이션 도구로는 오픈소스 보행자 시뮬레이션 라이브러리인 JuPedSim을 사용한다.'
    )

    add_styled_table(doc,
        ['항목', '내용'],
        [
            ['시뮬레이션 도구', 'JuPedSim (Python, 오픈소스)'],
            ['보행 모델', 'Collision-Free Speed Model V2 (Tordeux et al., 2016)'],
            ['게이트 선택 모델', 'MNL — Gao et al. (2019) LRP 기반'],
            ['대상역', '성수역 2호선 서쪽 대합실 (50m x 25m, 게이트 7개)'],
            ['시뮬레이션 시간', '300초 (열차 1~2회 도착 관찰)'],
            ['V&V 프레임워크', 'NIST TN 1822 (Ronchi et al., 2013) 준용'],
        ],
        col_widths=[5, 12]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-1. ').bold = True
    p.add_run('시뮬레이션 프레임워크 구성')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # 3.2 보행 모델: CFSM V2
    # =========================================================================
    doc.add_heading('3.2 보행 모델: CFSM V2', level=2)

    doc.add_heading('3.2.1 모델 개요', level=3)
    doc.add_paragraph(
        'CFSM V2 (Collision-Free Speed Model V2)는 Tordeux et al. (2016)이 제안한 '
        '1차 속도 기반 보행자 역학 모델이다. 각 보행자 i의 이동은 다음 수식으로 기술된다:'
    )

    # 수식
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq.add_run('x\u0307\u1d62 = V(s\u1d62) \u00b7 e\u1d62')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph(
        '여기서 V(s)는 최적 속도 함수로, 전방 보행자와의 간격 s에 따라 속도를 결정한다:'
    )

    p_eq2 = doc.add_paragraph()
    p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq2.add_run('V(s) = min{ v\u2080, max{ 0, (s - l) / T } }')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph(
        '여기서 v\u2080는 희망속도, l은 보행자 직경, T는 시간 간격 파라미터이다. '
        '이 모델의 핵심 특성은 간격 s가 보행자 직경 l 이하로 줄어들면 속도가 정확히 0에 수렴하여 '
        '보행자 간 겹침(overlap)이 구조적으로 발생하지 않는다는 점이다.'
    )

    doc.add_heading('3.2.2 모델 선택 근거: GCFM에서 CFSM V2로의 전환', level=3)
    doc.add_paragraph(
        '본 연구는 초기에 GCFM (Generalized Centrifugal Force Model, Chraibi et al., 2010)을 '
        '보행 모델로 채택하였으나, 다음과 같은 구조적 문제가 확인되어 CFSM V2로 전환하였다.'
    )

    add_styled_table(doc,
        ['문제', 'GCFM', 'CFSM V2'],
        [
            ['진동 (oscillation)', '힘 기반(F=ma) 2차 모델로\n좁은 게이트(0.55m)에서\n반발력-구동력 경합 → 진동', '1차 속도 모델로\n진동 원천 불가'],
            ['겹침 (overlap)', '반발력 상한(3.0N) 초과 시\n보행자 간 관통 발생\n(V6: 최대 22.4cm)', '충돌 방지 수학적 보장\n(V6: 겹침 0건)'],
            ['dt 민감도', '힘 적분 정밀도에 의존\n(V5: 24% 변동, FAIL)', '힘 적분 없음\n구조적으로 안정'],
            ['계산 효율', 'dt ≤ 0.01s 필수\n2차 적분 → 느림', 'dt = 0.05s 가능\n약 10배 빠름'],
        ],
        col_widths=[3.5, 6.5, 6.5]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-2. ').bold = True
    p.add_run('GCFM과 CFSM V2 비교')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.2.3 CFSM V2 파라미터', level=3)
    doc.add_paragraph(
        'CFSM V2의 파라미터는 Tordeux et al. (2016) 원논문 기본값을 바탕으로 하되, '
        'time_gap은 Seyfried et al. (2009)의 병목 유량 실험 데이터에 맞추어 캘리브레이션하였다.'
    )

    add_styled_table(doc,
        ['파라미터', '값', '단위', '출처/근거'],
        [
            ['time_gap', '0.80', 's', 'Seyfried (2009) 병목 유량 캘리브레이션'],
            ['radius', '0.15', 'm', '0.55m 게이트 통로 대응'],
            ['strength_neighbor_repulsion', '8.0', '-', 'Tordeux et al. (2016)'],
            ['range_neighbor_repulsion', '0.1', 'm', 'Tordeux et al. (2016)'],
            ['strength_geometry_repulsion', '5.0', '-', 'Tordeux et al. (2016)'],
            ['range_geometry_repulsion', '0.02', 'm', 'Tordeux et al. (2016)'],
            ['희망속도', 'N(1.34, 0.26)\nclip [0.8, 1.5]', 'm/s', 'Weidmann (1993)'],
            ['dt', '0.01', 's', '수치 안정성'],
        ],
        col_widths=[5.5, 3.5, 2, 6]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-3. ').bold = True
    p.add_run('CFSM V2 파라미터')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # 3.3 게이트 선택 모델
    # =========================================================================
    doc.add_heading('3.3 게이트 선택 모델: MNL (Gao LRP)', level=2)

    doc.add_paragraph(
        '보행자의 게이트 선택 행태는 Gao et al. (2019)의 LRP (Logit-based Route choice for '
        'Pedestrians) 모델을 기반으로 한다. 각 보행자는 게이트까지의 보행시간과 대기시간을 '
        '고려하여 효용을 계산하고, 다항 로짓(MNL) 확률로 게이트를 선택한다.'
    )

    doc.add_heading('3.3.1 효용함수', level=3)

    p_eq3 = doc.add_paragraph()
    p_eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_eq3.add_run('V\u2c7c = \u03c9\u1d3a \u00b7 W\u2c7c + \u03c9\u1d38 \u00b7 (L\u2081\u2c7c + L\u2083\u2c7c) / v\u1d62')
    run.font.size = Pt(11)
    run.italic = True

    doc.add_paragraph(
        '여기서 W\u2c7c는 게이트 j의 예상 대기시간, L\u2081\u2c7c는 현재 위치에서 게이트 j까지의 거리, '
        'L\u2083\u2c7c는 게이트 통과 후 출구까지의 거리, v\u1d62는 보행자 i의 희망속도이다.'
    )

    doc.add_heading('3.3.2 보행자 성격 유형', level=3)
    doc.add_paragraph(
        'Gao et al. (2019)에 따라 보행자를 세 가지 성격 유형으로 분류하며, '
        '각 유형은 대기시간과 보행시간에 대한 가중치가 다르다.'
    )

    add_styled_table(doc,
        ['성격 유형', '\u03c9\u1d3a (대기시간)', '\u03c9\u1d38 (보행시간)', '행태 특성'],
        [
            ['adventurous', '1.2', '0.8', '대기 회피, 먼 게이트도 감수'],
            ['conserved', '0.8', '1.2', '가까운 게이트 선호'],
            ['mild', '1.0', '1.0', '균등 고려'],
        ],
        col_widths=[3.5, 3.5, 3.5, 6]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-4. ').bold = True
    p.add_run('보행자 성격 유형별 가중치')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3.3 3단계 의사결정', level=3)
    doc.add_paragraph(
        '보행자는 게이트에 접근하면서 3단계에 걸쳐 게이트 선택을 수정한다.'
    )

    add_styled_table(doc,
        ['단계', '거리', '방법', '설명'],
        [
            ['1차 선택', '3.0m', 'MNL 확률적', 'Influence Zone 진입, 전체 게이트 고려'],
            ['2차 재선택', '1.7m', 'MNL 확률적', '접근거리 기반 재평가'],
            ['3차 재선택', '1.0m', '확정적 전환', '인접 빈 게이트로 즉시 전환'],
        ],
        col_widths=[3, 2.5, 3, 8]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-5. ').bold = True
    p.add_run('게이트 선택 3단계 의사결정')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # 3.4 서비스 시간 모델
    # =========================================================================
    doc.add_heading('3.4 서비스 시간 모델', level=2)

    doc.add_paragraph(
        '게이트 통과 시 소요되는 서비스 시간은 이용자 유형(태그/태그리스)에 따라 다르게 모델링한다.'
    )

    add_styled_table(doc,
        ['이용자 유형', '서비스 시간', '통과 방식', '출처'],
        [
            ['태그 (NFC)', 'lognormal\n평균 2.0s, [0.8, 3.7s]', '정지 → 태핑(1.1s)\n→ 문 열림 → 통과(0.65m/s)', 'Gao (2019) 실측'],
            ['태그리스', '0s (가정)', '감속 없이 보행속도로\n연속 통과', '우이신설선 실측 예정'],
        ],
        col_widths=[3.5, 4, 5, 4]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-6. ').bold = True
    p.add_run('이용자 유형별 서비스 시간 모델')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '태그리스 서비스 시간은 현재 0초로 가정하고 있으나, 우이신설선 현장조사를 통해 '
        '실측값으로 대체할 예정이다. 실측이 어려울 경우 0s, 0.5s, 1.0s 세 가지 시나리오로 '
        '민감도 분석을 수행한다.'
    )

    # =========================================================================
    # 3.5 V&V 프레임워크
    # =========================================================================
    doc.add_heading('3.5 검증 및 타당성 확인 (V&V)', level=2)

    doc.add_paragraph(
        '본 연구는 NIST TN 1822 (Ronchi et al., 2013)의 V&V 프레임워크를 준용하여 '
        '시뮬레이션 모델의 신뢰성을 체계적으로 확보한다. 기존 보행 시뮬레이션 연구의 약 80%가 '
        '체계적 V&V 없이 수행되고 있어 (Lovreglio et al., 2015), 본 연구에서는 6단계 '
        'V&V 절차를 따른다.'
    )

    add_styled_table(doc,
        ['Phase', '이름', '핵심 질문', '세부 항목 수'],
        [
            ['1', 'Model Qualification', '왜 CFSM V2 + MNL인가?', '5'],
            ['2', 'Verification', '코드가 수학 모델과 일치하는가?', '7 (V1~V7)'],
            ['3', 'Calibration', '실측에 맞게 파라미터를 조정', '6 (C1~C6)'],
            ['4', 'Validation', '독립 데이터로 현실과 비교', '5 (VD1~VD5)'],
            ['5', 'Uncertainty & Sensitivity', '결과가 안정적인가?', '3 (U1~U3)'],
            ['6', 'Application', '시나리오 실행 및 결과 보고', '1'],
        ],
        col_widths=[1.5, 4.5, 5.5, 3]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-7. ').bold = True
    p.add_run('V&V 프레임워크 6단계 구조 (NIST TN 1822 준용)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3.5.1 Verification ---
    doc.add_heading('3.5.1 Phase 2: Verification (검증)', level=3)

    doc.add_paragraph(
        'Verification 단계에서는 구현된 시뮬레이션 코드가 의도한 수학적 모델과 일치하는지 확인한다. '
        'RiMEA v4.1.1 가이드라인 및 NIST TN 1822 §3.1에 따라 7개 검증 테스트를 수행한다.'
    )

    add_styled_table(doc,
        ['ID', '테스트', '방법', '통과 기준', '근거'],
        [
            ['V1', '자유보행 속도', '30m x 5m 복도\n단일 보행자, 6개 희망속도', '실측/희망 오차 < 5%', 'RiMEA Test 1'],
            ['V2', '기본 다이어그램', '50m x 4m 복도\n밀도 0.5~4.0 P/m²', 'RMSE < 0.20 m/s\n단조감소', 'RiMEA Test 4\nWeidmann (1993)'],
            ['V3', '병목 유량', '10m x 8m 방 → 병목 → 방\n100명, 폭 0.6/0.8/1.0m', 'Seyfried (2009)\n대비 10% 이내', 'Seyfried et al.\n(2009)'],
            ['V4', '게이트 선택 규칙', 'MNL 1000회 반복\n거리/대기열/성격 효과', 'χ² 검정 p > 0.05\n성격 순서 일치', 'Gao (2019)'],
            ['V5', '수치 수렴성', 'dt = 0.01 / 0.005\n각 3회 반복', '평균 비유량\n변동 < 5%', 'NIST TN 1822\n§3.1.2'],
            ['V6', '물리 일관성', '밀도 1.5 P/m² 복도\n2s warmup 후 겹침 측정', '겹침 0건\n(Level A)', 'Chraibi (2010)\n§3'],
            ['V7', '대칭 테스트', '좌우 대칭 배치\n20회 반복', '좌우 비율 차이\n< 5%', 'RiMEA Test 6'],
        ],
        col_widths=[1, 2.5, 4, 3, 3]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-8. ').bold = True
    p.add_run('Verification 테스트 설계 (V1~V7)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3.5.2 Calibration ---
    doc.add_heading('3.5.2 Phase 3: Calibration (캘리브레이션)', level=3)

    doc.add_paragraph(
        'Calibration 단계에서는 실측 데이터에 맞게 모델 파라미터를 조정한다. '
        '캘리브레이션 데이터와 Validation 데이터는 반드시 분리하여 사용한다 (NIST TN 1822 §3.2).'
    )

    add_styled_table(doc,
        ['데이터 출처', '캘리브레이션용', 'Validation용'],
        [
            ['우이신설선 피크 촬영', 'O (서비스 시간, 접근속도)', 'X'],
            ['우이신설선 비피크 촬영', 'X', 'O (독립 검증)'],
            ['Seyfried (2009) 실험', 'O (병목 유량)', 'X'],
            ['Gao (2019) 실측', 'O (서비스 시간 분포)', 'X'],
            ['게이트별 도착률 (피크)', 'O', 'X'],
            ['게이트별 도착률 (비피크)', 'X', 'O'],
        ],
        col_widths=[5.5, 5.5, 5.5]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-9. ').bold = True
    p.add_run('데이터 분할 전략')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc,
        ['ID', '파라미터', '캘리브레이션 방법', '목표'],
        [
            ['C1', 'CFSM V2 파라미터\n(time_gap 등)', '병목 유량이\nSeyfried ±10% 이내', 'Js 오차 < 10%'],
            ['C2', '태그 서비스 시간 분포', '우이신설선 실측\n→ lognormal 피팅', 'KS 검정 p > 0.05'],
            ['C3', '태그리스 서비스 시간', '우이신설선 실측\n→ 분포 피팅 또는 상수', '실측값 반영'],
            ['C4', 'MNL 가중치\n(ω_N, ω_L)', '게이트별 도착률 관측\n→ 가중치 역추정', '게이트 분배\n오차 < 15%'],
            ['C5', '접근/퇴출 보행속도', '우이신설선 실측\n→ 희망속도 분포 업데이트', '분포 업데이트'],
            ['C6', '보행자 성격 비율', '직접 측정 불가\n→ 1:1:1 유지', '민감도 분석\n으로 보완'],
        ],
        col_widths=[1, 4, 5, 4]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-10. ').bold = True
    p.add_run('캘리브레이션 항목 (C1~C6)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3.5.3 Validation ---
    doc.add_heading('3.5.3 Phase 4: Validation (타당성 검증)', level=3)

    doc.add_paragraph(
        'Validation 단계에서는 캘리브레이션에 사용하지 않은 독립 데이터로 모델과 현실을 비교한다.'
    )

    add_styled_table(doc,
        ['ID', '테스트', '비교 대상', '통과 기준'],
        [
            ['VD1', '병목 유량\n(독립 폭)', 'Seyfried (2009)\n미사용 폭 (0.7/0.9/1.2m)', '오차 < 15%'],
            ['VD2', '대기행렬 길이', '우이신설선 비피크 실측', '평균 오차 < 30%\n또는 순위 일치'],
            ['VD3', '게이트별\n분배 비율', '우이신설선 비피크\n5분 단위 관측', '게이트별 비율\n오차 < 15%'],
            ['VD4', '서비스 시간\n분포 재현', '시뮬레이션 vs 실측\n서비스 시간 분포', 'KS 검정\np > 0.05'],
            ['VD5', 'Face Validation', '시뮬레이션 애니메이션을\n전문가에게 제시', '치명적 지적 없음'],
        ],
        col_widths=[1.5, 3, 5, 4]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-11. ').bold = True
    p.add_run('Validation 항목 (VD1~VD5)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3.5.4 Uncertainty ---
    doc.add_heading('3.5.4 Phase 5: 불확실성 및 민감도 분석', level=3)

    doc.add_paragraph(
        'NIST TN 1822 §4에 따라 시뮬레이션 결과의 불확실성을 정량화하고, '
        '핵심 결론의 강건성을 확인한다.'
    )

    add_styled_table(doc,
        ['ID', '분석', '방법', '통과 기준'],
        [
            ['U1', '반복 수렴성', '동일 시나리오 N회 반복\n누적 평균 수렴 확인', 'ERD < 5%\n연속 3회 유지'],
            ['U2', '파라미터\n민감도', 'OAT (One-at-a-time)\n±20% 변동, 9개 파라미터', '주요 파라미터 순위\n모델 결론 불변'],
            ['U3', '시나리오\n강건성', '모든 OAT 조건에서\n기본 vs 개선 비교', '결론 뒤집힘\n0건'],
        ],
        col_widths=[1, 3, 5, 4]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-12. ').bold = True
    p.add_run('불확실성 및 민감도 분석 (U1~U3)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'U2 민감도 분석의 대상 파라미터는 다음과 같다: '
        'time_gap, strength_neighbor_repulsion, strength_geometry_repulsion, '
        '태그 서비스 시간 평균, 태그리스 서비스 시간, MNL ω_N, MNL ω_L, '
        '보행자 성격 비율, 열차 하차 인원. '
        '각 파라미터를 기준값의 ±20% 범위에서 변동시키며 결과 변화를 관찰한다.'
    )

    # =========================================================================
    # 3.6 시나리오 설계
    # =========================================================================
    doc.add_heading('3.6 시나리오 설계', level=2)

    doc.add_heading('3.6.1 기본 시나리오 (현행 — 겸용)', level=3)
    doc.add_paragraph(
        '현행 운영 방식을 재현한다. 게이트 7개 중 1~2개를 태그/태그리스 겸용으로 운영하고, '
        '나머지는 태그 전용으로 운영한다. 겸용 게이트에서는 태그 이용자의 정지가 뒤쪽 태그리스 '
        '이용자의 대기를 유발하여 비효율이 발생한다.'
    )

    doc.add_heading('3.6.2 개선 시나리오 (전용 분리)', level=3)
    doc.add_paragraph(
        '겸용 게이트를 태그리스 전용으로 전환한다. 태그리스 이용자는 전용 게이트로만 진입하여 '
        '연속 통과가 가능하고, 태그 이용자는 나머지 게이트를 이용한다.'
    )

    doc.add_heading('3.6.3 비교 변수', level=3)

    add_styled_table(doc,
        ['변수', '수준', '값'],
        [
            ['배치 방식', '2', '겸용 / 전용 분리'],
            ['태그리스 이용자 비율 (p)', '5', '10%, 20%, 40%, 60%, 80%'],
            ['태그리스 전용 게이트 수 (k)', '3', '1대, 2대, 3대'],
            ['시간대', '2', '피크 / 비피크'],
        ],
        col_widths=[5, 2, 8]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-13. ').bold = True
    p.add_run('시나리오 비교 변수')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        '총 시나리오 수는 2(배치) x 5(비율) x 3(게이트 수) x 2(시간대) = 60개이며, '
        '각 시나리오를 Phase 5(U1)에서 도출한 최소 반복 수 N회 이상 반복하여 '
        '통계적 유의성을 확보한다.'
    )

    # =========================================================================
    # 3.7 평가 지표
    # =========================================================================
    doc.add_heading('3.7 평가 지표', level=2)

    add_styled_table(doc,
        ['지표', '정의', '단위'],
        [
            ['평균 대기시간', '게이트 도착 ~ 서비스 시작 시간 평균', '초'],
            ['최대 대기행렬 길이', '게이트 앞 대기 인원 최대값', '명'],
            ['총 통과시간', '계단 출발 ~ 게이트 통과 완료', '초'],
            ['게이트 균형도 (MD)', '게이트별 이용 편차', '%'],
            ['보행자 밀도 / LOS', '게이트-계단 구간 밀도', '인/m²'],
            ['처리량', '단위시간당 통과 인원', 'person/min'],
        ],
        col_widths=[4, 8, 3]
    )

    p = doc.add_paragraph()
    p.add_run('표 3-14. ').bold = True
    p.add_run('시뮬레이션 평가 지표')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # 저장
    # =========================================================================
    output_path = OUTPUT_DIR / '3장_방법론.docx'
    doc.save(str(output_path))
    print(f"저장 완료: {output_path}")


if __name__ == '__main__':
    create_report()
